from io import BytesIO

import pytest
from docx import Document

from src import document_io
from src.document_io import extract_text


def test_extracts_utf8_text_file():
    assert extract_text("consent.txt", "개인정보 동의".encode()) == "개인정보 동의"


def test_extracts_docx_paragraphs():
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("첫 번째 문장")
    document.add_paragraph("두 번째 문장")
    document.save(buffer)

    text = extract_text("consent.docx", buffer.getvalue())

    assert "첫 번째 문장" in text
    assert "두 번째 문장" in text


def test_rejects_unsupported_file_type():
    with pytest.raises(ValueError):
        extract_text("consent.xlsx", b"invalid")


def test_image_uses_ocr_extractor(monkeypatch):
    monkeypatch.setattr(
        document_io,
        "_extract_ocr_text",
        lambda content, extension: "OCR 계약서 내용",
    )

    assert extract_text("contract.png", b"image") == "OCR 계약서 내용"
